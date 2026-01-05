from pathlib import Path
from datetime import datetime
from typing import Optional, List
from sqlalchemy.orm import Session
import re

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models import Property, Lead, LetterTemplate, LetterHistory
from ..config import settings


class LetterService:
    """Service for generating letters from templates"""
    
    DEFAULT_TEMPLATES = [
        {
            "name": "Initial Offer Letter",
            "subject": "Cash Offer for {{address}}",
            "body": """Dear {{owner_name}},

I am writing to express interest in purchasing your property located at:

{{address}}
{{city}}, FL {{zip}}
Folio #: {{folio_number}}

I am a local investor looking to purchase properties in Broward County for cash, with a quick closing timeline.

If you're interested in selling, please contact me at your earliest convenience.

Best regards,
[Your Name]
[Your Phone]
[Your Email]""",
            "template_type": "mail"
        },
        {
            "name": "Follow-Up Letter",
            "subject": "Following Up - {{address}}",
            "body": """Dear {{owner_name}},

I recently reached out regarding your property at {{address}}.

I wanted to follow up and see if you've had a chance to consider my offer. I remain interested in purchasing your property for cash.

Property Details:
- Address: {{address}}, {{city}}, FL {{zip}}
- Assessed Value: ${{just_value}}

Please feel free to reach out with any questions.

Best regards,
[Your Name]""",
            "template_type": "mail"
        },
        {
            "name": "Absentee Owner Letter",
            "subject": "Interested in Your Broward County Property",
            "body": """Dear {{owner_name}},

I noticed you own a property in Broward County, FL:

{{address}}
{{city}}, FL {{zip}}

As a local real estate investor, I specialize in purchasing properties from out-of-state owners who may be interested in selling.

Benefits of working with me:
• Cash purchase - no financing contingencies
• Quick closing - as fast as 7-14 days
• Buy "as-is" - no repairs needed
• I cover all closing costs

If you've ever considered selling, I'd love to discuss options with you.

Best regards,
[Your Name]""",
            "template_type": "mail"
        },
        {
            "name": "High Equity Opportunity",
            "subject": "Quick Cash Offer for {{address}}",
            "body": """Dear {{owner_name}},

I hope this letter finds you well. I am a real estate investor actively purchasing properties in your area.

I am specifically interested in your property:
{{address}}
{{city}}, FL {{zip}}

Based on my research, your property shows strong potential equity, which makes it a great candidate for a quick cash sale.

Current Market Indicators:
- Assessed Value: ${{just_value}}
- Estimated Purchase Price: ${{estimated_price}}
- Potential Equity: ${{potential_equity}}

I can close quickly, often within 2 weeks, and purchase the property in its current condition.

If you're interested in a no-obligation cash offer, please contact me.

Best regards,
[Your Name]""",
            "template_type": "mail"
        }
    ]
    
    @classmethod
    def init_default_templates(cls, db: Session):
        """Initialize database with default letter templates"""
        existing = db.query(LetterTemplate).count()
        if existing == 0:
            for template_data in cls.DEFAULT_TEMPLATES:
                template = LetterTemplate(**template_data)
                db.add(template)
            db.commit()
    
    @classmethod
    def get_templates(cls, db: Session) -> List[LetterTemplate]:
        """Get all letter templates"""
        return db.query(LetterTemplate).all()
    
    @classmethod
    def get_template(cls, db: Session, template_id: int) -> Optional[LetterTemplate]:
        """Get a specific template"""
        return db.query(LetterTemplate).filter(LetterTemplate.id == template_id).first()
    
    @classmethod
    def create_template(cls, db: Session, data: dict) -> LetterTemplate:
        """Create a new template"""
        template = LetterTemplate(**data)
        db.add(template)
        db.commit()
        db.refresh(template)
        return template
    
    @classmethod
    def update_template(cls, db: Session, template_id: int, data: dict) -> Optional[LetterTemplate]:
        """Update an existing template"""
        template = cls.get_template(db, template_id)
        if template:
            for key, value in data.items():
                setattr(template, key, value)
            db.commit()
            db.refresh(template)
        return template
    
    @classmethod
    def _get_template_variables(cls, property: Property, lead: Optional[Lead] = None) -> dict:
        """Build template variable dictionary from property data"""
        address = ' '.join(filter(None, [
            property.situs_street_number,
            property.situs_street_name,
            property.situs_street_type
        ]))
        
        mailing_address = ' '.join(filter(None, [
            property.mailing_address_line_1,
            property.mailing_address_line_2
        ]))
        
        mailing_city_state_zip = ', '.join(filter(None, [
            property.mailing_city,
            property.mailing_state
        ]))
        if property.mailing_zip:
            mailing_city_state_zip += f' {property.mailing_zip}'
        
        def format_currency(value):
            if value is None:
                return 'N/A'
            return f'{value:,.0f}'
        
        return {
            'owner_name': property.name_line_1 or 'Property Owner',
            'owner_name_2': property.name_line_2 or '',
            'address': address,
            'city': property.situs_city or '',
            'zip': property.situs_zip or '',
            'folio_number': property.folio_number,
            'just_value': format_currency(property.just_value),
            'estimated_price': format_currency(property.estimated_purchase_price),
            'potential_equity': format_currency(property.potential_equity),
            'mailing_address': mailing_address,
            'mailing_address_line_1': property.mailing_address_line_1 or '',
            'mailing_address_line_2': property.mailing_address_line_2 or '',
            'mailing_city': property.mailing_city or '',
            'mailing_state': property.mailing_state or '',
            'mailing_zip': property.mailing_zip or '',
            'mailing_city_state_zip': mailing_city_state_zip,
            'use_type': property.use_type or '',
            'year_built': str(property.bldg_year_built or ''),
            'sq_footage': format_currency(property.bldg_tot_sq_footage) if property.bldg_tot_sq_footage else '',
            'beds': str(property.beds or ''),
            'baths': str(property.baths or ''),
            'date': datetime.now().strftime('%B %d, %Y'),
            # Lead fields
            'lead_status': lead.lead_status if lead else '',
            'notes': lead.notes if lead else '',
        }
    
    @classmethod
    def _replace_variables(cls, text: str, variables: dict) -> str:
        """Replace template variables in text"""
        for key, value in variables.items():
            text = text.replace(f'{{{{{key}}}}}', str(value))
        return text
    
    @classmethod
    def generate_letter(
        cls,
        db: Session,
        folio_number: str,
        template_id: int,
        output_format: str = 'pdf',
        output_path: Optional[Path] = None
    ) -> dict:
        """
        Generate a letter for a property using a template
        
        Args:
            db: Database session
            folio_number: Property folio number
            template_id: Letter template ID
            output_format: 'pdf' or 'docx'
            output_path: Optional custom output path
            
        Returns:
            dict with success status and file path
        """
        # Get property
        property = db.query(Property).filter(Property.folio_number == folio_number).first()
        if not property:
            raise ValueError(f"Property not found: {folio_number}")
        
        # Get lead if exists
        lead = db.query(Lead).filter(Lead.folio_number == folio_number).first()
        
        # Get template
        template = cls.get_template(db, template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")
        
        # Build variables and replace
        variables = cls._get_template_variables(property, lead)
        subject = cls._replace_variables(template.subject or '', variables)
        body = cls._replace_variables(template.body or '', variables)
        
        # Generate output path if not provided
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_folio = re.sub(r'[^a-zA-Z0-9]', '_', folio_number)
            filename = f"letter_{safe_folio}_{timestamp}.{output_format}"
            output_path = settings.LETTERS_DIR / filename
        
        output_path = Path(output_path)
        
        # Generate file
        if output_format == 'pdf':
            cls._generate_pdf(output_path, subject, body, variables)
        else:
            cls._generate_docx(output_path, subject, body, variables)
        
        # Log to history
        history = LetterHistory(
            folio_number=folio_number,
            template_id=template_id,
            file_path=str(output_path)
        )
        db.add(history)
        db.commit()
        
        return {
            'success': True,
            'file_path': str(output_path),
            'filename': output_path.name
        }
    
    @classmethod
    def _generate_pdf(cls, output_path: Path, subject: str, body: str, variables: dict):
        """Generate PDF letter"""
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=20
        )
        
        address_style = ParagraphStyle(
            'AddressStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6
        )
        
        subject_style = ParagraphStyle(
            'SubjectStyle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            spaceBefore=20,
            spaceAfter=20
        )
        
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            alignment=TA_LEFT
        )
        
        # Build content
        content = []
        
        # Date
        content.append(Paragraph(variables['date'], date_style))
        
        # Recipient address
        if variables.get('mailing_address_line_1'):
            content.append(Paragraph(variables['owner_name'], address_style))
            content.append(Paragraph(variables['mailing_address_line_1'], address_style))
            if variables.get('mailing_address_line_2'):
                content.append(Paragraph(variables['mailing_address_line_2'], address_style))
            content.append(Paragraph(variables['mailing_city_state_zip'], address_style))
        
        # Subject
        if subject:
            content.append(Paragraph(f"Re: {subject}", subject_style))
        
        content.append(Spacer(1, 12))
        
        # Body - convert newlines to HTML breaks
        body_html = body.replace('\n\n', '</para><para>').replace('\n', '<br/>')
        for para in body_html.split('</para><para>'):
            content.append(Paragraph(para, body_style))
            content.append(Spacer(1, 12))
        
        doc.build(content)
    
    @classmethod
    def _generate_docx(cls, output_path: Path, subject: str, body: str, variables: dict):
        """Generate DOCX letter"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Date
        date_para = doc.add_paragraph(variables['date'])
        date_para.space_after = Pt(24)
        
        # Recipient address
        if variables.get('mailing_address_line_1'):
            doc.add_paragraph(variables['owner_name'])
            doc.add_paragraph(variables['mailing_address_line_1'])
            if variables.get('mailing_address_line_2'):
                doc.add_paragraph(variables['mailing_address_line_2'])
            doc.add_paragraph(variables['mailing_city_state_zip'])
        
        # Subject
        if subject:
            subject_para = doc.add_paragraph()
            subject_para.space_before = Pt(24)
            subject_run = subject_para.add_run(f"Re: {subject}")
            subject_run.bold = True
        
        doc.add_paragraph()  # Spacer
        
        # Body
        for line in body.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()  # Empty paragraph for spacing
        
        doc.save(str(output_path))
    
    @classmethod
    def generate_bulk_letters(
        cls,
        db: Session,
        folio_numbers: List[str],
        template_id: int,
        output_format: str = 'pdf'
    ) -> dict:
        """Generate letters for multiple properties"""
        results = []
        success_count = 0
        error_count = 0
        
        for folio in folio_numbers:
            try:
                result = cls.generate_letter(db, folio, template_id, output_format)
                results.append({
                    'folio_number': folio,
                    'success': True,
                    'file_path': result['file_path']
                })
                success_count += 1
            except Exception as e:
                results.append({
                    'folio_number': folio,
                    'success': False,
                    'error': str(e)
                })
                error_count += 1
        
        return {
            'total': len(folio_numbers),
            'success_count': success_count,
            'error_count': error_count,
            'results': results,
            'output_directory': str(settings.LETTERS_DIR)
        }



